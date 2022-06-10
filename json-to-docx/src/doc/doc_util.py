#!/usr/bin/env python3
'''
various utilities for formatting a docx
'''

import sys
import lxml
import random
import string
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.table import _Cell


if sys.platform == 'win32':
	import win32com.client as client

from helper.logger import *


# --------------------------------------------------------------------------------------------------------------------------------------------
# pictures, background image

''' graphic-style
'''
def create_graphic_style(doc, valign, halign):
    style_name = f"fr-{random_string()}"

    # graphic_properties_attributes = {'wrap': 'none', 'verticalpos': valign, 'horizontalpos': halign}
    # graphic_properties = style.GraphicProperties(attributes=graphic_properties_attributes)
    #
    # graphic_style_attributes = {'name': style_name, 'family': 'graphic', 'parentstylename': 'Graphics'}
    # graphic_style = style.Style(attributes=graphic_style_attributes)
    #
    # graphic_style.addElement(graphic_properties)
    # odt.automaticstyles.addElement(graphic_style)

    return style_name


''' frame and image
'''
def create_image_frame(doc, picture_path, valign, halign, width, height):
    # THIS IS THE Draw:Frame object to return
    draw_frame = None

    # first the image to be added into the document
    # href = doc.addPicture(picture_path)
    # if href:
    #     # next we need the Draw:Image object
    #     image_attributes = {'href': href}
    #     # image_attributes[('draw', 'mimetype')] = 'image/png'
    #     draw_image = draw.Image(attributes=image_attributes)
    #
    #     frame_style_name = create_graphic_style(doc, valign, halign)
    #
    #     # finally we need the Draw:Frame object
    #     frame_attributes = {'stylename': frame_style_name, 'anchortype': 'paragraph', 'width': f"{width}in", 'height': f"{height}in"}
    #     draw_frame = draw.Frame(attributes=frame_attributes)
    #
    #     draw_frame.addElement(draw_image)
    #
    # else:
    #     warn(f"image {picture_path} copuld not be added into the document")

    return draw_frame



# --------------------------------------------------------------------------------------------------------------------------------------------
# table, table-row, table-column, table-cell

''' create a Table
'''
def create_table(doc, table_name, table_style_attributes, table_properties_attributes):
    tbl = None

    if 'family' not in table_style_attributes:
        table_style_attributes['family'] = 'table'

    # create the style
    # table_style = style.Style(attributes=table_style_attributes)
    # table_style.addElement(style.TableProperties(attributes=table_properties_attributes))
    # odt.automaticstyles.addElement(table_style)

    # create the table
    # table_properties = {'name': table_name, 'stylename': table_style_attributes['name']}
    # tbl = table.Table(attributes=table_properties)

    return tbl


''' create table-header-rows
'''
def create_table_header_rows():
    return table.TableHeaderRows()


''' create TableColumn
'''
def create_table_column(doc, table_column_name, table_column_style_attributes, table_column_properties_attributes):
    table_column = None

    if 'family' not in table_column_style_attributes:
        table_column_style_attributes['family'] = 'table-column'

    # create the style
    # table_column_style = style.Style(attributes=table_column_style_attributes)
    # table_column_style.addElement(style.TableColumnProperties(attributes=table_column_properties_attributes))
    # odt.automaticstyles.addElement(table_column_style)
    #
    # # create the table-column
    # table_column_properties = {'stylename': table_column_style_attributes['name']}
    # table_column = table.TableColumn(attributes=table_column_properties)

    return table_column


''' create TableRow
'''
def create_table_row(doc, table_row_style_attributes, table_row_properties_attributes):
    table_row = None

    if 'family' not in table_row_style_attributes:
        table_row_style_attributes['family'] = 'table-row'

    # create the style
    # table_row_style = style.Style(attributes=table_row_style_attributes)
    # table_row_properties_attributes['keeptogether'] = 'always'
    # table_row_style.addElement(style.TableRowProperties(attributes=table_row_properties_attributes))
    # odt.automaticstyles.addElement(table_row_style)
    #
    # # create the table-row
    # table_row_properties = {'stylename': table_row_style_attributes['name']}
    # table_row = table.TableRow(attributes=table_row_properties)

    return table_row


''' create TableCell
'''
def create_table_cell(doc, table_cell_style_attributes, table_cell_properties_attributes, table_cell_attributes, background_image_style=None):
    table_cell = None

    if 'family' not in table_cell_style_attributes:
        table_cell_style_attributes['family'] = 'table-cell'

    # create the style
    # table_cell_style = style.Style(attributes=table_cell_style_attributes)
    # table_cell_properties = style.TableCellProperties(attributes=table_cell_properties_attributes)
    #
    # if background_image_style:
    #     table_cell_properties.addElement(background_image_style)
    #
    # table_cell_style.addElement(table_cell_properties)
    # odt.automaticstyles.addElement(table_cell_style)
    #
    # # create the table-cell
    # table_cell_attributes['stylename'] = table_cell_style_attributes['name']
    # table_cell = table.TableCell(attributes=table_cell_attributes)

    return table_cell


''' create CoveredTableCell
'''
def create_covered_table_cell(doc, table_cell_style_attributes, table_cell_properties_attributes):
    table_cell = None

    if 'family' not in table_cell_style_attributes:
        table_cell_style_attributes['family'] = 'table-cell'

    # create the style
    # table_cell_style = style.Style(attributes=table_cell_style_attributes)
    # table_cell_style.addElement(style.TableCellProperties(attributes=table_cell_properties_attributes))
    # odt.automaticstyles.addElement(table_cell_style)
    #
    # # create the table-cell
    # table_cell_attributes = {'stylename': table_cell_style_attributes['name']}
    # table_cell = table.CoveredTableCell(attributes=table_cell_attributes)

    return table_cell



# --------------------------------------------------------------------------------------------------------------------------------------------
# paragraphs and styles

''' create style - family paragraph
'''
def create_paragraph_style(doc, style_attributes=None, paragraph_attributes=None, text_attributes=None):
    # we may need to create the style-name
    if style_attributes is None:
        style_attributes = {}

    if 'family' not in style_attributes:
        style_attributes['family'] = 'paragraph'

    if 'name' not in style_attributes:
        style_attributes['name'] = random_string()

    if 'parentstylename' not in style_attributes:
        style_attributes['parentstylename'] = 'Text_20_body'


    paragraph_style = None

    # create the style
    # paragraph_style = style.Style(attributes=style_attributes)
    #
    # if paragraph_attributes is not None:
    #     paragraph_style.addElement(style.ParagraphProperties(attributes=paragraph_attributes))
    #
    # if text_attributes is not None:
    #     paragraph_style.addElement(style.TextProperties(attributes=text_attributes))
    #
    # odt.automaticstyles.addElement(paragraph_style)

    return paragraph_style.getAttribute('name')


''' page number
'''
def create_page_number(style_name, short=False):
    paragraph = text.P(stylename=style_name)
    page_number = text.PageNumber(selectpage='current')

    if short:
        paragraph.addText("Page ")
        paragraph.addElement(page_number)
    else:
        s = text.S()
        page_count = text.PageCount()

        paragraph.addText("Page ")
        paragraph.addElement(page_number)
        paragraph.addElement(s)
        paragraph.addText("of ")
        paragraph.addElement(page_count)

    return paragraph


''' write a paragraph in a given style
'''
def create_paragraph(doc, style_name, text_content=None, run_list=None, outline_level=0):
    paragraph = None

    # style = odt.getStyleByName(style_name)
    # if style is None:
    #     warn(f"style {style_name} not found")
    #
    # if run_list is not None:
    #     paragraph = text.P(stylename=style)
    #     for run in run_list:
    #         style_attributes = {'family': 'text'}
    #         text_style_name = create_paragraph_style(doc, style_attributes=style_attributes, text_attributes=run['text-attributes'])
    #         paragraph.addElement(text.Span(stylename=text_style_name, text=run['text']))
    #
    #     return paragraph
    #
    # if text_content is not None:
    #     if outline_level == 0:
    #         paragraph = text.P(stylename=style_name, text=text_content)
    #     else:
    #         paragraph = text.H(stylename=style_name, text=text_content, outlinelevel=outline_level)
    #
    #     return paragraph
    # else:
    #     paragraph = text.P(stylename=style_name)
    #     return paragraph

    return paragraph

# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# indexes and pdf generation

''' update docx indexes by opening and closing the docx, rest is done by macros
'''
def update_indexes(docx_path):
	try:
		word = client.DispatchEx("Word.Application")
		doc = word.Documents.Open(docx_path)
		doc.Close()
	except Exception as e:
		raise e
	finally:
		word.Quit()


''' set docx updateFields property true
'''
def set_updatefields_true(docx_path):
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    doc = Document(docx_path)
    # add child to doc.settings element
    element_updatefields = lxml.etree.SubElement(
        doc.settings.element, f"{namespace}updateFields"
    )
    element_updatefields.set(f"{namespace}val", "true")
    doc.save(docx_path)


''' given an docx file generates pdf in the given directory
'''
def generate_pdf(infile, outdir):
	pdf_path = infile.replace(".docx", r".pdf")
	try:
		word = client.DispatchEx("Word.Application")
		doc = word.Documents.Open(infile)
		doc.SaveAs(pdf_path, FileFormat = 17)
		doc.Close()
	except Exception as e:
		raise e
	finally:
		word.Quit()


''' create table-of-contents
'''
def create_toc():
    name = 'Table of Content'
    toc = text.TableOfContent(name=name)
    toc_source = text.TableOfContentSource(outlinelevel=10)
    toc_title_template = text.IndexTitleTemplate()
    toc_source.addElement(toc_title_template)

    for i in range(1, 11):
        toc_entry_template = text.TableOfContentEntryTemplate(outlinelevel=i, stylename=f"Contents_20_{i}")

        index_entry_link_start = text.IndexEntryLinkStart(stylename='Index_20_Link')
        index_entry_chapter = text.IndexEntryChapter()
        index_entry_text = text.IndexEntryText()
        index_entry_tab_stop = text.IndexEntryTabStop(type='right', leaderchar='.')
        index_entry_page_number = text.IndexEntryPageNumber()
        index_entry_link_end = text.IndexEntryLinkEnd()

        toc_entry_template.addElement(index_entry_link_start)
        toc_entry_template.addElement(index_entry_chapter)
        toc_entry_template.addElement(index_entry_text)
        toc_entry_template.addElement(index_entry_tab_stop)
        toc_entry_template.addElement(index_entry_page_number)
        toc_entry_template.addElement(index_entry_link_end)

        toc_source.addElement(toc_entry_template)

    toc.addElement(toc_source)

    return toc


''' create illustration-index
'''
def create_lof():
    name = 'List of Figures'
    toc = text.TableOfContent(name=name)
    toc_source = text.TableOfContentSource(outlinelevel=1, useoutlinelevel=False, useindexmarks=False, useindexsourcestyles=True)
    toc_entry_template = text.TableOfContentEntryTemplate(outlinelevel=1, stylename='Figure')

    index_entry_link_start = text.IndexEntryLinkStart(stylename='Index_20_Link')
    index_entry_chapter = text.IndexEntryChapter()
    index_entry_text = text.IndexEntryText()
    index_entry_tab_stop = text.IndexEntryTabStop(type='right', leaderchar='.')
    index_entry_page_number = text.IndexEntryPageNumber()
    index_entry_link_end = text.IndexEntryLinkEnd()

    toc_entry_template.addElement(index_entry_link_start)
    toc_entry_template.addElement(index_entry_chapter)
    toc_entry_template.addElement(index_entry_text)
    toc_entry_template.addElement(index_entry_tab_stop)
    toc_entry_template.addElement(index_entry_page_number)
    toc_entry_template.addElement(index_entry_link_end)

    toc_title_template = text.IndexTitleTemplate()
    toc_index_source_styles = text.IndexSourceStyles(outlinelevel=1)
    toc_index_source_style = text.IndexSourceStyle(stylename='Figure')
    toc_index_source_styles.addElement(toc_index_source_style)

    toc_source.addElement(toc_entry_template)
    toc_source.addElement(toc_title_template)
    toc_source.addElement(toc_index_source_styles)
    toc.addElement(toc_source)

    return toc


''' create Table-index
'''
def create_lot():
    name = 'List of Tables'
    toc = text.TableOfContent(name=name)
    toc_source = text.TableOfContentSource(outlinelevel=1, useoutlinelevel=False, useindexmarks=False, useindexsourcestyles=True)
    toc_entry_template = text.TableOfContentEntryTemplate(outlinelevel=1, stylename='Table')

    index_entry_link_start = text.IndexEntryLinkStart(stylename='Index_20_Link')
    index_entry_chapter = text.IndexEntryChapter()
    index_entry_text = text.IndexEntryText()
    index_entry_tab_stop = text.IndexEntryTabStop(type='right', leaderchar='.')
    index_entry_page_number = text.IndexEntryPageNumber()
    index_entry_link_end = text.IndexEntryLinkEnd()

    toc_entry_template.addElement(index_entry_link_start)
    toc_entry_template.addElement(index_entry_chapter)
    toc_entry_template.addElement(index_entry_text)
    toc_entry_template.addElement(index_entry_tab_stop)
    toc_entry_template.addElement(index_entry_page_number)
    toc_entry_template.addElement(index_entry_link_end)

    toc_title_template = text.IndexTitleTemplate()
    toc_index_source_styles = text.IndexSourceStyles(outlinelevel=1)
    toc_index_source_style = text.IndexSourceStyle(stylename='Table')
    toc_index_source_styles.addElement(toc_index_source_style)

    toc_source.addElement(toc_entry_template)
    toc_source.addElement(toc_title_template)
    toc_source.addElement(toc_index_source_styles)
    toc.addElement(toc_source)

    return toc



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# master-page and page-layout

''' get master-page by name
'''
def get_master_page(doc, master_page_name):
    # for master_page in odt.masterstyles.getElementsByType(style.MasterPage):
    #     if master_page.getAttribute('name') == master_page_name:
    #         return master_page

    warn(f"master-page {master_page_name} NOT found")
    return None


''' get page-layout by name
'''
def get_page_layout(doc, page_layout_name):
    # for page_layout in odt.automaticstyles.getElementsByType(style.PageLayout):
    #     if page_layout.getAttribute('name') == page_layout_name:
    #         return page_layout

    warn(f"page-layout {page_layout_name} NOT found")
    return None


''' update page-layout of Standard master-page with the given page-layout
'''
def update_master_page_page_layout(doc, master_page_name, new_page_layout_name):
    master_page = get_master_page(doc, master_page_name)

    if master_page is not None:
        master_page.attributes[(master_page.qname[0], 'page-layout-name')] = new_page_layout_name


''' create (section-specific) page-layout
'''
def create_page_layout(doc, docx_specs, page_layout_name, page_spec, margin_spec, orientation):
    page_layout = None
    # create one
    # page_layout = style.PageLayout(name=page_layout_name)
    # odt.automaticstyles.addElement(page_layout)

    if orientation == 'portrait':
        pageheight = f"{docx_specs['page-spec'][page_spec]['height']}in"
        pagewidth = f"{docx_specs['page-spec'][page_spec]['width']}in"
    else:
        pageheight = f"{docx_specs['page-spec'][page_spec]['width']}in"
        pagewidth = f"{docx_specs['page-spec'][page_spec]['height']}in"

    margintop = f"{docx_specs['margin-spec'][margin_spec]['top']}in"
    marginbottom = f"{docx_specs['margin-spec'][margin_spec]['bottom']}in"
    marginleft = f"{docx_specs['margin-spec'][margin_spec]['left']}in"
    marginright = f"{docx_specs['margin-spec'][margin_spec]['right']}in"
    # margingutter = f"{docx_specs['margin-spec'][margin_spec]['gutter']}in"

    # page_layout.addElement(style.PageLayoutProperties(pagewidth=pagewidth, pageheight=pageheight, margintop=margintop, marginbottom=marginbottom, marginleft=marginleft, marginright=marginright, printorientation=orientation))

    return page_layout


''' create (section-specific) master-page
    page layouts are saved with a name mp-section-no
'''
def create_master_page(doc, docx_specs, master_page_name, page_layout_name, page_spec, margin_spec, orientation):
    master_page = None

    # create one, first get/create the page-layout
    # page_layout = create_page_layout(doc, docx_specs, page_layout_name, page_spec, margin_spec, orientation)
    # master_page = style.MasterPage(name=master_page_name, pagelayoutname=page_layout_name)
    # odt.masterstyles.addElement(master_page)

    return master_page


''' create header/footer
'''
def create_header_footer(master_page, page_layout, header_footer, odd_even):
    header_footer_style = None
    if header_footer == 'header':
        height = f"{HEADER_HEIGHT}in"
        if odd_even == 'odd':
            header_footer_style = style.Header()
        if odd_even == 'even':
            header_footer_style = style.HeaderLeft()
        if odd_even == 'first':
            # header_footer_style = HeaderFirst()
            pass

    elif header_footer == 'footer':
        height = f"{FOOTER_HEIGHT}in"
        if odd_even == 'odd':
            header_footer_style = style.Footer()
        if odd_even == 'even':
            header_footer_style = style.FooterLeft()
        if odd_even == 'first':
            # header_footer_style = FooterFirst()
            pass

    if header_footer_style:
        # TODO: the height should come from actual header content height
        header_footer_properties_attributes = {'margin': '0in', 'padding': '0in', 'dynamicspacing': False, 'height': height}
        header_style = style.HeaderStyle()
        header_style.addElement(style.HeaderFooterProperties(attributes=header_footer_properties_attributes))
        footer_style = style.FooterStyle()
        footer_style.addElement(style.HeaderFooterProperties(attributes=header_footer_properties_attributes))
        page_layout.addElement(header_style)
        page_layout.addElement(footer_style)

        master_page.addElement(header_footer_style)

    return header_footer_style



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# various utility functions

''' whether the container is a table-cell or not
'''
def is_table_cell(container):
    # if container is n instance of table-cell
    if isinstance(container, type(table.TableCell)):
        return True
    else:
        return False


''' given pixel size, calculate the row height in inches
    a reasonable approximation is what gsheet says 21 pixels, renders well as 12 pixel (assuming our normal text is 10-11 in size)
'''
def row_height_in_inches(pixel_size):
    return float((pixel_size) / 96)


''' get a random string
'''
def random_string(length=12):
    letters = string.ascii_uppercase
    return ''.join(random.choice(letters) for i in range(length))


''' fit width/height into a given width/height maintaining aspect ratio
'''
def fit_width_height(fit_within_width, fit_within_height, width_to_fit, height_to_fit):
    WIDTH_OFFSET = 0.0
    HEIGHT_OFFSET = 0.2

    fit_within_width = fit_within_width - WIDTH_OFFSET
    fit_within_height = fit_within_height - HEIGHT_OFFSET

    aspect_ratio = width_to_fit / height_to_fit

    if width_to_fit > fit_within_width:
        width_to_fit = fit_within_width
        height_to_fit = width_to_fit / aspect_ratio
        if height_to_fit > fit_within_height:
            height_to_fit = fit_within_height
            width_to_fit = height_to_fit * aspect_ratio

    return width_to_fit, height_to_fit



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# various utility data

COLSEP = (0/72)
# ROWSEP = (2/72)

HEADER_HEIGHT = 0.3
FOOTER_HEIGHT = 0.3

GSHEET_ODT_BORDER_MAPPING = {
    'DOTTED': 'dotted',
    'DASHED': 'dash',
    'SOLID': 'solid'
}


HEADING_TO_LEVEL = {
    'Heading 1': {'outline-level': 1},
    'Heading 2': {'outline-level': 2},
    'Heading 3': {'outline-level': 3},
    'Heading 4': {'outline-level': 4},
    'Heading 5': {'outline-level': 5},
    'Heading 6': {'outline-level': 6},
    'Heading 7': {'outline-level': 7},
    'Heading 8': {'outline-level': 8},
    'Heading 9': {'outline-level': 9},
    'Heading 10': {'outline-level': 10},
}


LEVEL_TO_HEADING = [
    'Title',
    'Heading 1',
    'Heading 2',
    'Heading 3',
    'Heading 4',
    'Heading 5',
    'Heading 6',
    'Heading 7',
    'Heading 8',
    'Heading 9',
    'Heading 10',
]


TEXT_VALIGN_MAP = {'TOP': 'top', 'MIDDLE': 'middle', 'BOTTOM': 'bottom'}
# TEXT_HALIGN_MAP = {'LEFT': 'start', 'CENTER': 'center', 'RIGHT': 'end', 'JUSTIFY': 'justify'}
TEXT_HALIGN_MAP = {'LEFT': 'left', 'CENTER': 'center', 'RIGHT': 'right', 'JUSTIFY': 'justify'}

IMAGE_POSITION = {'center': 'center', 'middle': 'center', 'left': 'left', 'right': 'right', 'top': 'top', 'bottom': 'bottom'}

WRAP_STRATEGY_MAP = {'OVERFLOW': 'no-wrap', 'CLIP': 'no-wrap', 'WRAP': 'wrap'}

COLUMNS = [ 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z',
            'AA', 'AB', 'AC', 'AD', 'AE', 'AF', 'AG', 'AH', 'AI', 'AJ', 'AK', 'AL', 'AM', 'AN', 'AO', 'AP', 'AQ', 'AR', 'AS', 'AT', 'AU', 'AV', 'AW', 'AX', 'AY', 'AZ',
            'BA', 'BB', 'BC', 'BD', 'BE', 'BF', 'BG', 'BH', 'BI', 'BJ', 'BK', 'BL', 'BM', 'BN', 'BO', 'BP', 'BQ', 'BR', 'BS', 'BT', 'BU', 'BV', 'BW', 'BX', 'BY', 'BZ']




GSHEET_OXML_BORDER_MAPPING = {
    'DOTTED': 'dotted',
    'DASHED': 'dashed',
    'SOLID': 'single',
    'SOLID_MEDIUM': 'thick',
    'SOLID_THICK': 'triple',
    'DOUBLE': 'double',
    'NONE': 'none'
}

'''
Table of Contents
'''
def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    # creates a new element
    fldChar = OxmlElement('w:fldChar')
    # sets attribute on element
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    # sets attribute on element
    instrText.set(qn('xml:space'), 'preserve')
    # change 1-3 depending on heading levels you need
    instrText.text = 'TOC \\o "1-6" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update Table of Contents."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

'''
List of Figures
'''
def add_lof(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    # creates a new element
    fldChar = OxmlElement('w:fldChar')
    # sets attribute on element
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    # sets attribute on element
    instrText.set(qn('xml:space'), 'preserve')
    # change 1-3 depending on heading levels you need
    instrText.text = 'TOC \\h \\z \\t "Figure" \\c'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update List of Figures."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

'''
List of Tables
'''
def add_lot(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    # creates a new element
    fldChar = OxmlElement('w:fldChar')
    # sets attribute on element
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    # sets attribute on element
    instrText.set(qn('xml:space'), 'preserve')
    # change 1-3 depending on heading levels you need
    instrText.text = 'TOC \\h \\z \\t "Table" \\c'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update List of Figures."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

def add_numbered_paragraph(doc, text, restart=False, style='List Number'):
    if restart:
        # new numbering
        ABSTRACT_NUM_ID = 8
        numbering = doc._part.numbering_part.numbering_definitions._numbering
        num_id = numbering._next_numId
        num = CT_Num.new(num_id, ABSTRACT_NUM_ID)
        num.add_lvlOverride(ilvl=0).add_startOverride(1)
        w_num = numbering._insert_num(num)

        #w_num_pr_xml = '<w:ilvl w:val="0"/><w:numId w:val="{0}"/>'.format(num_id)
        #w_num_pr = CT_NumPr()
        p = doc.add_paragraph(text, style)

         # Access paragraph XML element
        p_xml = p._p

        # Paragraph properties
        p_props = p_xml.get_or_add_pPr()

        # Create number properties element
        num_props = OxmlElement('w:numPr')

        lvl_prop = OxmlElement('w:ilvl')
        lvl_prop.set(qn('w:val'), '0')

        num_id_prop = OxmlElement('w:numId')
        num_id_prop.set(qn('w:val'), str(w_num.numId))

        num_props.append(lvl_prop)
        num_props.append(num_id_prop)

        # Add number properties to paragraph
        p_props.append(num_props)

    else:
        doc.add_paragraph(text, style)

def add_horizontal_line(paragraph, pos='w:bottom', size='6', color='auto'):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement(pos)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)

def append_page_number_only(paragraph):
    run = paragraph.add_run()

    # page number
    # creates a new element
    fldCharBegin1 = OxmlElement('w:fldChar')
    # sets attribute on element
    fldCharBegin1.set(qn('w:fldCharType'), 'begin')
    instrText1 = OxmlElement('w:instrText')
    # sets attribute on element
    instrText1.set(qn('xml:space'), 'preserve')
    # page number
    instrText1.text = 'PAGE \* MERGEFORMAT'

    fldCharSeparate1 = OxmlElement('w:fldChar')
    fldCharSeparate1.set(qn('w:fldCharType'), 'separate')

    fldCharEnd1 = OxmlElement('w:fldChar')
    fldCharEnd1.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldCharBegin1)
    r_element.append(instrText1)
    r_element.append(fldCharSeparate1)
    r_element.append(fldCharEnd1)
    p_element = paragraph._p

def append_page_number_with_pages(paragraph, separator=' of '):
    run = paragraph.add_run()

    # page number
    # creates a new element
    fldCharBegin1 = OxmlElement('w:fldChar')
    # sets attribute on element
    fldCharBegin1.set(qn('w:fldCharType'), 'begin')
    instrText1 = OxmlElement('w:instrText')
    # sets attribute on element
    instrText1.set(qn('xml:space'), 'preserve')
    # page number
    instrText1.text = 'PAGE \* MERGEFORMAT'

    fldCharSeparate1 = OxmlElement('w:fldChar')
    fldCharSeparate1.set(qn('w:fldCharType'), 'separate')

    fldCharEnd1 = OxmlElement('w:fldChar')
    fldCharEnd1.set(qn('w:fldCharType'), 'end')

    fldCharOf = OxmlElement('w:t')
    fldCharOf.set(qn('xml:space'), 'preserve')
    fldCharOf.text = separator

    # number of pages
    # creates a new element
    fldCharBegin2 = OxmlElement('w:fldChar')
    # sets attribute on element
    fldCharBegin2.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    # sets attribute on element
    instrText2.set(qn('xml:space'), 'preserve')
    # page number
    instrText2.text = 'NUMPAGES \* MERGEFORMAT'

    fldCharSeparate2 = OxmlElement('w:fldChar')
    fldCharSeparate2.set(qn('w:fldCharType'), 'separate')

    fldCharEnd2 = OxmlElement('w:fldChar')
    fldCharEnd2.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldCharBegin1)
    r_element.append(instrText1)
    r_element.append(fldCharSeparate1)
    r_element.append(fldCharEnd1)
    r_element.append(fldCharOf)
    r_element.append(fldCharBegin2)
    r_element.append(instrText2)
    r_element.append(fldCharSeparate2)
    r_element.append(fldCharEnd2)
    p_element = paragraph._p

def rotate_text(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)

def set_character_style(run, spec):
    run.bold = spec['bold']
    run.italic = spec['italic']
    run.strike = spec['strikethrough']
    run.underline = spec['underline']

    run.font.name = spec['fontFamily']
    run.font.size = Pt(spec['fontSize'])

    red, green, blue = 0, 0, 0
    fgcolor = spec['foregroundColor']
    if fgcolor == {}:
        red, green, blue = 32, 32, 32
    else:
        red = int(fgcolor['red'] * 255) if 'red' in fgcolor else 0
        green = int(fgcolor['green'] * 255) if 'green' in fgcolor else 0
        blue = int(fgcolor['blue'] * 255) if 'blue' in fgcolor else 0

    run.font.color.rgb = RGBColor(red, green, blue)

def set_cell_bgcolor(cell, color):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)

def set_paragraph_bgcolor(paragraph, color):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    paragraph._p.get_or_add_pPr().append(shading_elm_1)

def copy_cell_border(from_cell: _Cell, to_cell: _Cell):
    from_tc = from_cell._tc
    from_tcPr = from_tc.get_or_add_tcPr()

    to_tc = to_cell._tc
    to_tcPr = to_tc.get_or_add_tcPr()

    from_tcBorders = from_tcPr.first_child_found_in("w:tcBorders")
    to_tcBorders = to_tcPr.first_child_found_in("w:tcBorders")
    if from_tcBorders is not None:
        if to_tcBorders is None:
            to_tcBorders = deepcopy(from_tcBorders)
            to_tc.get_or_add_tcPr().append(to_tcBorders)

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell's border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_paragraph_border(paragraph, **kwargs):
    """
    Set paragraph's border
    Usage:

    set_paragraph_border(
        paragraph,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    pPr = paragraph._p.get_or_add_pPr()

    # check for tag existnace, if none found, then create one
    pBorders = pPr.first_child_found_in("w:pBorders")
    if pBorders is None:
        pBorders = OxmlElement('w:pBorders')
        pPr.append(pBorders)

    # list over all available tags
    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_str = edge
            if edge_str == 'start': edge_str = 'left'
            if edge_str == 'end': edge_str = 'right'
            tag = 'w:{}'.format(edge_str)

            # check for tag existnace, if none found, then create one
            element = pBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                pBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def merge_document(placeholder, docx_path):
    # the document is in the same folder as the template, get the path
    # docx_path = os.path.abspath('{0}/{1}'.format('../conf', docx_name))
    sub_doc = Document(docx_path)

    par_parent = placeholder._p.getparent()
    index = par_parent.index(placeholder._p) + 1
    for element in sub_doc.part.element:
        element.remove_all('w:sectPr')
        par_parent.insert(index, element)
        index = index + 1

def polish_table(table):
    for r in table.rows:
        # no preferred width for the last column
        c = r._tr.tc_lst[-1]
        #for c in r._tr.tc_lst:
        tcW = c.tcPr.tcW
        tcW.type = 'auto'
        tcW.w = 0

def ooxml_border_from_gsheet_border(borders, key):
    if key in borders:
        border = borders[key]
        red = int(border['color']['red'] * 255) if 'red' in border['color'] else 0
        green = int(border['color']['green'] * 255) if 'green' in border['color'] else 0
        blue = int(border['color']['blue'] * 255) if 'blue' in border['color'] else 0
        color = '{0:02x}{1:02x}{2:02x}'.format(red, green, blue)
        if 'style' in border:
            border_style = border['style']
        else:
            border_style = 'NONE'

        return {"sz": border['width'] * 8, "val": GSHEET_OXML_BORDER_MAPPING[border_style], "color": color, "space": "0"}
    else:
        return None

def insert_image(cell, image_spec):
    '''
        image_spec is like {'url': url, 'path': local_path, 'height': height, 'width': width, 'dpi': im_dpi}
    '''
    if image_spec is not None:
        run = cell.paragraphs[0].add_run()
        run.add_picture(image_spec['path'], height=Pt(image_spec['height']), width=Pt(image_spec['width']))

def set_repeat_table_header(row):
    ''' set repeat table row on every new page
    '''
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row
